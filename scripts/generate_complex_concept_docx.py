#!/usr/bin/env python3
"""生成视频素材：非常复杂的长篇模组构想 Word 文档（约 13000 字）。"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn

OUTPUT = Path(__file__).resolve().parent.parent / "OutputFootage" / "complex-mod-concept.docx"

TITLE = "《裂隙编年史：跨维度公会领土、动态经济与社会模拟》模组完整设计规格书"
SUBTITLE = "Minecraft 1.20.1 · Fabric · 客户端与服务端双端必装 · 设计稿 v0.9.3-draft"
AUTHOR = "构想作者：视频演示用虚构案例 · 非最终产品文档"


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def add_para(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r2 = p.add_run(text)
    r2.font.name = "宋体"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def build_sections() -> list[tuple[str, int, list[str]]]:
    """返回 (章节标题, heading_level, 段落列表)。"""
    sections: list[tuple[str, int, list[str]]] = []

    sections.append(
        (
            "0. 文档说明与范围界定",
            1,
            [
                "本文档描述一款面向中大型私服与整合包作者的 Fabric 模组完整构想，目标是在不破坏原版生存节奏的前提下，引入可配置的跨维度领土、公会社会结构、动态 NPC 经济、以及基于「裂隙能量」的资源循环体系。文档面向策划、程序、美术与 QA 共同评审，所有机制均需在单机、专用服务器、以及 Paper 兼容环境下可降级运行。",
                "非目标（Out of Scope）包括：完全替代现有知名领地插件/模组的全部功能、官方 Marketplace 商业化、基岩版移植、以及需要修改 Minecraft 核心协议的大规模网络同步改造。首版 MVP 必须可独立游玩，后续版本再开放与 Vault、LuckPerms、Dynmap 等的可选桥接。",
                "读者须知：文中出现的数值均为占位默认值，最终数值需通过 common 配置文件与 datapack 覆盖；所有「必须/应当/可以」措辞遵循 RFC 2119 语义，便于 Agent 或人工开发时做优先级裁剪。",
            ],
        )
    )

    sections.append(
        (
            "1. 核心愿景与体验支柱",
            1,
            [
                "体验支柱一：「有意义的领土」。玩家建立的不仅是方块边界，而是具有历史叙事、资源产出、法律规则与外交关系的聚落实体。领土应随时间演化——从营地、哨站、城镇到要塞都市，每个阶段解锁不同建筑蓝图、NPC 职业与裂隙科技。",
                "体验支柱二：「社会即玩法」。公会不是聊天频道附加物，而是资源分配、战争动员、立法与贸易的执行主体。玩家可以扮演执政官、大法官、军需官、外交官等角色，各角色拥有不同的 UI 面板与操作权限，且权限可被公会长自定义。",
                "体验支柱三：「经济有摩擦」。完全自由的玩家间交易容易导致通胀与脚本刷钱。本模组引入 NPC 商行、动态供需、运输成本、关税与仓储损耗，使「倒卖」仍有利可图但需承担风险与时间成本。",
                "体验支柱四：「跨维度同一套规则」。主世界、下界、末地以及模组自定义维度「裂隙荒原」使用统一的领土与外交协议，但各维度可有差异化参数（如下界领土维护费加倍、裂隙荒原允许 PVP 常开）。",
            ],
        )
    )

    sections.append(
        (
            "1.1 目标玩家画像",
            2,
            [
                "Primary：20–50 人长期开服的服主与核心玩家，已有基础插件/模组管理经验，希望用一套深度系统替代零散插件堆叠。",
                "Secondary：整合包作者，需要可配置关闭 PVP/战争模块，仅保留经济与社会模拟作为 RPG 背景。",
                "Tertiary：单机玩家，通过「顾问 NPC」与可调难度参数，以较低管理负担体验城镇经营与裂隙探索。",
            ],
        )
    )

    # 大量机制细节章节
    mechanics = [
        (
            "2. 裂隙能量与世界观底层",
            [
                "裂隙能量（Rift Energy, RE）是本模组的核心抽象资源，不可直接放入玩家背包，以「公国账本」「区块蓄能器」「玩家灵魂绑定的谐振值」三种形态存在。RE 的产出与消耗必须形成闭环，避免无上限囤积。",
                "世界观设定（可配置关闭显示）：古代文明在维度褶皱处留下裂隙锚点，现代玩家可通过谐振装置唤醒锚点，获得科技树解锁与领土合法宣称权。叙事文本通过进度系统与结构物调查笔记发放，不强制打断 gameplay。",
                "RE 产出源：裂隙锚点被动滴答（受区块加载与在线玩家数影响）、特定怪物掉落谐振碎片、完成公会里程碑、以及与其他公会的贸易顺差。",
                "RE 消耗项：领土维护（按区块数指数级软上限）、高级建筑激活、跨维度传送门维持、战争动员（宣战后的 72 游戏小时内持续 burn）、以及「法令」施行的行政成本。",
            ],
        ),
        (
            "2.1 裂隙荒原维度",
            [
                "新增维度「裂隙荒原」，入口通过主世界特定 Y 层与 RE 阈值触发，或通过建造「界门框架」稳定开启。维度地形算法需与主世界种子关联但使用独立噪声叠加，保证可复现。",
                "荒原特性：随机「法则扰动」事件（重力变化、怪物强化、资源富集窗口），每 3–7 游戏日轮换；死亡惩罚可配置为 RE 损失、物品掉落保护、或传统掉落。",
                "荒原领土规则：可宣称但维护费 ×1.8；允许其他公会「劫掠」资源点，劫掠成功获得 RE 与稀有材料，失败则宣战冷却延长。",
            ],
        ),
        (
            "3. 公会系统与社会等级",
            [
                "公会创建需：① 三名玩家联名；② 支付 RE 注册费；③ 在主世界放置「誓约碑」并完成命名（过滤敏感词，支持管理员自定义黑名单）。",
                "等级体系：游荡者 → 成员 → 精英 → 官员 → 副长 → 公会长。每级默认权限矩阵可在 guild 配置中覆盖，且支持「自定义职位」——最多 12 个，每个职位绑定权限位掩码与工资模板。",
                "公会声望：通过完成合同、抵御袭击、学术捐献（向 NPC 大学捐赠物品）提升。声望解锁外交选项（互不侵犯、关税同盟、军事同盟、联邦合并）。",
                "解散与继承：公会长离线超过配置天数触发「摄政」机制；长期无维护则领土进入「荒废」状态，可被其他公会低价并购。",
            ],
        ),
        (
            "3.1 立法与法令系统",
            [
                "公会长与拥有立法权限的官员可颁布「法令」，法令类型包括：税率调整、禁入区域、装备限制、PVP 开关、以及自定义脚本触发的轻量规则（如「日落之后禁止下界旅行」）。",
                "法令必须通过「议会投票」生效（可配置为独裁模式跳过）。投票期间成员可弃权；通过阈值默认 60% 参与成员赞成。",
                "跨公会条约：两项公会可签署条约，条约内容模板化 + 自由文本附录（附录长度限制 500 字，供 RP 使用）。违约检测由系统监控（如条约禁止攻击但发生伤害事件），违约方 RE 罚没与声望下降。",
            ],
        ),
        (
            "4. 领土宣称与区块管理",
            [
                "宣称单位：16×16 区块为最小单元，支持批量框选（Shift+工具）与多边形近似（顶点放置界桩）。最大宣称数随公会等级与 RE 蓄能器等级提升，硬顶由服主 config 控制。",
                "重叠处理：未宣称区域先到先得；重叠宣称进入「争议」状态，双方可通过外交、决斗裁决（可选迷你游戏）、或系统拍卖 RE 竞价解决。",
                "子区域：在一个宣称块内可划分「功能区」——工业（允许高频红石）、 residential（禁止 TNT）、 sacred（禁止挖掘）、 military（允许 PVP）。功能区边界用粒子显示，可仅对成员可见。",
                "访客权限：陌生人默认不可破坏/开箱/使用门；可通过「通行证」物品或临时授权（时长、范围可配）开放。",
            ],
        ),
        (
            "4.1 维护、荒废与并购",
            [
                "每 real-time 日（可改游戏日）扣除维护 RE。维护不足时依次触发：警告信鸽 NPC → 功能降级（禁止新建筑蓝图）→ 保护弱化（非成员可放置但不可破坏）→ 荒废（可被并购）。",
                "并购流程：收购方支付 RE + 游戏内货币（若启用桥接），被并购方公会长确认或超时默认拒绝。并购后誓约碑更换纹理，历史日志写入「编年史书」物品。",
            ],
        ),
        (
            "5. 动态经济与 NPC 商行",
            [
                "NPC 商行由「贸易协会」实体管理，在村庄、玩家建造的集市结构、以及裂隙荒原驿站生成。每个商行有独立库存、偏好商品、价格波动曲线与信用等级。",
                "价格模型：基础价 × 区域稀缺度 × 季节系数 × 最近 7 日交易量反馈。服主可锁定特定物品价格上下限，防止极端通胀/通缩。",
                "合同系统：公会可接取「运输合同」「清剿合同」「建设合同」，完成后获得 RE、声望与商行信誉。合同失败（超时、货物被劫）有惩罚。",
                "玩家店铺：允许在宣称区放置「柜台」方块，绑定库存与定价；支持离线交易（物品放入 escrow 容器）。",
            ],
        ),
        (
            "5.1 货币与通胀控制",
            [
                "默认不新增物理货币物品，使用 EMERALD 与 RE 双轨；整合包作者可配置自定义货币物品与兑换比率。",
                "系统定期回收：维修费、执照费、拍卖手续费、以及「奢侈税」（对超过阈值的单次交易抽成）。",
            ],
        ),
        (
            "6. 战争、袭击与冲突解析",
            [
                "宣战条件：双方公会等级 ≥2、无现有停战条约、发起方 RE ≥ 阈值。宣战后进入「战争状态」：特定维度 PVP 强制开启、领土维护费上升、但战争分数（War Score）通过占点、击杀、破坏军事建筑累积。",
                "战争分数达到 100 时触发「和谈窗口」，败方选择割让区块、赔款 RE、或停战 30 日。若 14 日内无和谈，系统按分数强制执行默认条款。",
                "非对称袭击：小股「裂隙劫掠者」会攻击维护不足的领土，强度与荒废度正相关，鼓励协防而非纯 PVP。",
                "友好火控：同公会/同盟成员间伤害默认关闭，可在战争区 override。",
            ],
        ),
        (
            "7. 建筑蓝图与结构物",
            [
                "蓝图系统：通过「建筑师工作台」消耗材料 + RE 逐块投影建造，支持暂停、撤销（限时）、以及多人协作加速（每人同时段有贡献上限防刷）。",
                "官方蓝图包：哨塔、市场、界门框架、蓄能器阵列、议会厅；整合包可通过 datapack 追加。",
                "结构检测：使用 jigsaw 或自定义模板匹配验证玩家是否完成关键结构以解锁科技节点。",
            ],
        ),
        (
            "8. UI、引导与无障碍",
            [
                "主 UI「编年史之书」：分页显示公会信息、领土地图（2D 俯视简化）、经济仪表盘、外交关系图、以及待办合同。",
                "新手引导：首次进服触发可跳过教程，分 5 步介绍宣称、RE、商行、法令、裂隙门。",
                "色盲模式：领土边界粒子提供形状差异；UI 图表避免纯红绿对比。",
                "多语言：至少 zh_cn、en_us；所有玩家可见文本进入 lang 文件，禁止硬编码。",
            ],
        ),
        (
            "9. 技术架构与模块边界",
            [
                "模块划分建议：core（事件总线、配置）、territory（宣称与权限）、guild（社会与外交）、economy（商行与合同）、rift（维度与 RE）、war（冲突）、ui（客户端屏幕）、compat（桥接）。",
                "服务端权威：所有领土变更、RE 账本、战争状态以服务端为准；客户端仅预测与渲染。",
                "数据持久化：SQLite 或 LevelDB 存储公会/领土/账本；支持 `/export` 与定时备份；单人存档与多人服使用相同 schema。",
                "网络包：自定义 payload 需版本化；大包（地图快照）分片传输，避免超过 Fabric 默认限制。",
                "性能预算：100 在线、5000 宣称区块时，tick 额外开销目标 < 2ms（服主可开「低精度模式」降低粒子与价格刷新频率）。",
            ],
        ),
        (
            "9.1 兼容性与可选集成",
            [
                "与 Fabric 常用模组：JEI/EMI 显示 RE 与合同；Waystones 界门可标记为「外交敏感」；FTB Chunks 并存时需 config 选择主领土系统。",
                "Optional API：暴露 TerritoryQuery、GuildMembership、REBalance 供附属模组调用；API  semver 从 1.0.0 起。",
            ],
        ),
        (
            "10. 配置项概览（common / client / server 分离）",
            [
                "common：启用模块开关、RE 全局倍率、维度规则、蓝图列表、语言默认。",
                "server：最大宣称、战争开关、维护周期、NPC 生成密度、敏感词、备份间隔。",
                "client：UI 缩放、粒子密度、地图风格、教程重置。",
                "所有配置变更需 hot-reload 或 `/riftconfig reload` 指令，并写 audit log。",
            ],
        ),
        (
            "11. 测试计划与验收标准",
            [
                "单元测试：价格公式、维护费曲线、权限矩阵、战争分数计算。",
                "集成测试：10 机器人客户端模拟宣称、交易、宣战、并购全流程。",
                " soak 测试：7 日不间断服，监控内存泄漏与区块卸载后数据一致性。",
                "验收：MVP 必须包含 RE 循环、基础宣称、公会创建、1 个商行、1 个裂隙锚点；战争与荒原维度可标记为 Beta。",
            ],
        ),
        (
            "12. 里程碑与版本规划",
            [
                "v0.1 MVP：RE + 宣称 + 公会 + 基础 UI。",
                "v0.3：商行 + 合同 + 维护荒废。",
                "v0.5：战争 + 条约 + 编年史日志。",
                "v0.8：裂隙荒原 + 蓝图系统。",
                "v1.0：API 稳定 + 文档 + 整合包预设配置包。",
            ],
        ),
        (
            "13. 开放问题（待规划 Agent 澄清）",
            [
                "是否在首版实现联邦合并，还是仅保留军事同盟？",
                "RE 是否允许玩家间直接转账，还是必须经公会账本？",
                "荒原维度是否与现有地形生成模组（Terralith 等）联动？",
                "单机模式下 NPC 顾问的 AI 对话深度：模板式 vs 可选 LLM 桥接（默认关闭）？",
                "反作弊：如何检测虚假客户端在宣称区快速放置？是否与常见反作弊模组接口？",
            ],
        ),
    ]

    for title, paras in mechanics:
        sections.append((title, 2 if title[0].isdigit() and title[1] == "." else 2, paras))

    # 附录：详细数值表（增加字数）
    appendices = [
        (
            "附录 A：默认数值表（占位）",
            1,
            [
                "RE 基础滴答：每宣称区块每小时 +0.05 RE（软上限前）；超过 200 区块后每新增区块效率 ×0.92。",
                "维护费：每区块每游戏日 0.12 RE；荒原 ×1.8；战争期间 ×1.3。",
                "宣战冷却：7 游戏日；停战后 14 游戏日不可对同一目标再次宣战。",
                "商行价格波动：单次交易影响 ±0.5%–3%，每日回归基准 15%。",
                "公会成员上限：1 级 8 人，每升 1 级 +4，硬顶 64（可 config）。",
            ],
        ),
        (
            "附录 B：关键物品与方块清单（摘要）",
            1,
            [
                "誓约碑、界桩、蓄能器核心、谐振碎片、编年史之书、建筑师工作台、柜台、通行证、战争旗帜、和约卷轴、裂隙罗盘。",
                "各物品需完整模型/纹理/loot table/进度触发；首版可用占位纹理但 ID 必须稳定。",
            ],
        ),
        (
            "附录 C：指令与权限节点",
            1,
            [
                "指令：`/guild`、`/claim`、`/rift`、`/treaty`、`/riftconfig`；子命令需补全与权限节点 `rift.admin`、`rift.guild.officer` 等。",
                "OP 与 Fabric 权限 API 双轨；LuckPerms 桥接为 optional 模块。",
            ],
        ),
        (
            "附录 D：术语表",
            1,
            [
                "RE（裂隙能量）、宣称（Claim）、荒废（Decay）、战争分数（War Score）、谐振（Attune）、界门（Rift Gate）、合同（Contract）、法令（Edict）、并购（Annexation）。",
            ],
        ),
    ]
    sections.extend(appendices)

    # 扩展段落以达到约 13000 字
    filler_topics = [
        (
            "14. 叙事事件与随机遭遇",
            2,
            [
                "除系统战争外，世界应定期生成「编年史事件」：流浪学者请求护送、裂隙兽潮预警、两个 NPC 商行价格战、古代锚点苏醒等。事件权重可配置，且需避免对小型服务器造成压倒性负担。",
                "每个事件有最低在线人数门槛与冷却；失败不应硬惩罚，但成功给予独特 cosmetic 或少量 RE。",
            ],
        ),
        (
            "14.1 季节与日历",
            2,
            [
                "可选「游戏日历」：每 28 游戏日为一个月，影响作物价格、怪物强度与节日（节日期间关税下调、特殊合同刷新）。",
                "日历与真实时间脱钩，服主可冻结或加速。",
            ],
        ),
        (
            "15. 音频与反馈",
            2,
            [
                "关键操作需音效反馈：宣称成功、RE 不足、宣战、条约签署、蓝图完成。音效包独立 resource location，支持资源包覆盖。",
                "背景音乐不强制；荒原维度可选环境音层叠。",
            ],
        ),
        (
            "16. 数据迁移与升级",
            2,
            [
                "版本升级时必须提供 NBT/SQL 迁移脚本；破坏性变更需 `--backup` 提示与 rollback 文档。",
                "从其他领土模组导入：仅提供实验性 CSV 导入，不保证完整映射。",
            ],
        ),
    ]
    sections.extend(filler_topics)

    # 大篇幅附录，确保 Word 打开后字数统计 ≥13000
    api_paras = []
    for i in range(1, 41):
        api_paras.append(
            f"API 端点草案 §{i}：TerritoryService.queryClaims(world, chunkX, chunkZ) 返回可选 ClaimView 列表，"
            f"含 ownerGuildId、subzoneType、maintenanceDebt、以及 flags 位集（PVP/BUILD/INTERACT/CHEST）。"
            f"第 {i} 版需保证查询在 5000 区块规模下 p95 延迟 < 8ms；若超出则自动降级为 chunk 粗粒度缓存。"
            f"所有 API 调用须通过服务端主线程或受控工作队列派发，禁止在渲染线程访问账本。"
        )
    sections.append(("22. 服务端 API 契约草案（节选）", 1, api_paras))

    return sections


def expand_to_target(doc: Document, sections: list, target_chars: int) -> int:
    """若不足目标字数，追加重复性详细说明（仍保持可读）。"""
    total = sum(len(p.text) for p in doc.paragraphs)
    if total >= target_chars:
        return total

    add_heading(doc, "17. 详细用例场景（User Stories）", level=1)
    stories = [
        "用例 US-001：三名玩家在新世界 24 小时内完成公会注册、首次宣称 4 区块、放置誓约碑，并收到新手合同「运送 64 原木至最近商行」。",
        "用例 US-002：公会 A 与 B 签署关税同盟后，A 的成员在 B 的商行购买物品享受 10% 折扣，折扣由双方 RE 池共同补贴。",
        "用例 US-003：维护不足 3 日的公会收到劫掠预警，10 名成员协作在 20 分钟内击退裂隙劫掠者，获得声望与谐振碎片。",
        "用例 US-004：整合包作者关闭战争模块，仅保留商行与蓝图，单机玩家通过顾问 NPC 完成「建设集市」进度线。",
        "用例 US-005：服主执行 `/riftconfig reload` 将最大宣称从 500 降至 300，系统按 LRU 标记超出部分进入 7 日宽限期。",
        "用例 US-006：玩家在荒原维度开启界门，遭遇法则扰动「低重力」，合同「采集 32 裂隙水晶」难度动态上调奖励。",
        "用例 US-007：公会长颁布法令「日落禁止下界」，系统通过维度 travel 事件拦截并提示；违规 3 次扣个人 RE。",
        "用例 US-008：并购流程中，被并购方公会长 48 小时未上线，副长可代确认；并购后编年史之书保留双方战争史摘要。",
    ]
    for s in stories:
        add_bullet(doc, s)
        total += len(s)

    detail_blocks = [
        (
            "18. 风险登记册",
            [
                "风险 R-01：大型服务器宣称数量导致查询延迟 —— 缓解：空间索引 + 异步缓存。",
                "风险 R-02：经济脚本刷合同 —— 缓解：合同贡献上限 + 异常检测。",
                "风险 R-03：与现有模组冲突 —— 缓解：compat 模块 + 已知不兼容列表文档。",
                "风险 R-04：玩家流失导致荒废地图 —— 缓解：并购与 NPC 回收机制。",
                "风险 R-05：翻译滞后 —— 缓解：en_us 与 zh_cn 同步 PR 检查。",
            ],
        ),
        (
            "19. 性能与监控指标",
            [
                "指标：每 tick 领土查询次数、RE 事务 TPS、网络包队列深度、SQLite 写入延迟。",
                "告警：维护批处理超过 50ms、单日 RE 净增发超过 config 阈值、战争状态 stuck 超过 30 日。",
                "服主面板：可选 Fabric 服务端 GUI 或 Web 仪表盘（Web 为远期，不在 MVP）。",
            ],
        ),
        (
            "20. 法务与内容安全",
            [
                "公会名、法令文本、条约附录需过敏感词与长度校验；支持管理员人工复核队列。",
                "玩家举报入口：编年史之书内一键提交最近 50 条相关日志片段。",
            ],
        ),
    ]

    for title, paras in detail_blocks:
        add_heading(doc, title, level=1)
        for para in paras:
            add_para(doc, para)
            total += len(para)

    # 最后一轮：补充 FAQ 直到达标
    if total < target_chars:
        add_heading(doc, "21. 常见问题 FAQ", level=1)
        faqs = [
            ("问：单机能否禁用 PVP 与战争？", "答：可以，server 配置一键关闭 war 模块，RE 循环仍可通过 PvE 与合同维持。"),
            ("问：是否必须安装 JEI？", "答：非必须，但强烈推荐；未安装时使用简化配方书。"),
            ("问：与原版村庄的关系？", "答：商行可绑定原版村庄中心，但不修改村民 AI 核心，仅追加交易选项与合同 NPC。"),
            ("问：数据包能否覆盖价格？", "答：可以，datapack 路径 `data/rift_chronicles/economy/`。"),
            ("问：如何降级卸载？", "答：卸载前导出 CSV；未宣称区域无残留，已宣称方块转为普通方块但丢失 RE 数据。"),
        ]
        for q, a in faqs:
            add_para(doc, a, bold_lead=q + " ")
            total += len(q) + len(a)

    # 若仍不足，添加细化段落
    extra_idx = 0
    while total < target_chars and extra_idx < 80:
        extra_idx += 1
        text = (
            f"补充说明 #{extra_idx}：在多人协作开发本模组时，第 {extra_idx} 批交付物应包含对应的集成测试、"
            f"配置样例与玩家可读补丁说明；任何涉及 RE 账本的全局变更必须写入不可篡改的 audit 日志，"
            f"以便服主在争议处理时追溯。开发优先级以 MVP 路径为准，非 MVP 功能可通过 `features.{extra_idx}` 开关禁用。"
        )
        add_para(doc, text)
        total += len(text)

    return total


def add_footer_word_count(doc: Document, char_count: int) -> None:
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"— 文档结束 —\n\n字数统计：{char_count}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(
        "（视频演示素材 · 用于表现「复杂构想」场景 · 非真实开发委托文档）"
    )
    nr.font.size = Pt(10)
    nr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    nr.font.name = "宋体"
    nr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def main() -> None:
    TARGET = 13000
    doc = Document()
    set_doc_defaults(doc)

    # 封面
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run(TITLE)
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.name = "黑体"
    tr._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    for line in (SUBTITLE, AUTHOR, ""):
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.add_page_break()

    sections = build_sections()
    for title, level, paras in sections:
        add_heading(doc, title, level=level)
        for para in paras:
            add_para(doc, para)

    char_count = expand_to_target(doc, sections, TARGET)

    char_count = sum(len(p.text) for p in doc.paragraphs)
    add_footer_word_count(doc, char_count)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    print(f"Character count (paragraph text): {char_count}")


if __name__ == "__main__":
    main()
